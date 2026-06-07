from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from analysis_pipeline import NHANESDepressionAnalysis
from generate_tp_manuscript import apply_fonts_to_run, apply_three_line_table_styling


ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "results"
OUT_MD = RESULTS / "supp_table7_missingness.md"
OUT_DOCX = RESULTS / "supp_table7_missingness.docx"
PACKAGE_MD = ROOT / "submission_packages" / "medrxiv" / "supp_table7_missingness.md"
PACKAGE_DOCX = ROOT / "submission_packages" / "medrxiv" / "supp_table7_missingness.docx"
TP_PACKAGE_MD = ROOT / "submission_packages" / "translational_psychiatry" / "supp_table7_missingness.md"
TP_PACKAGE_DOCX = ROOT / "submission_packages" / "translational_psychiatry" / "supp_table7_missingness.docx"


VARIABLES = [
    ("log_ratio", "log(GBR)", "Primary exposure; log(GGT/total bilirubin)"),
    ("LBXSGTSI", "Serum GGT", "Exposure component"),
    ("LBXSTB", "Total bilirubin", "Exposure component"),
    ("RIDAGEYR", "Age", "Covariate"),
    ("is_female", "Sex", "Covariate; female indicator"),
    ("BMXBMI", "BMI", "Covariate"),
    ("pir", "Poverty-to-income ratio", "Covariate"),
    ("education_cat", "Education category", "Covariate; missing category retained in models"),
    ("is_married", "Marital status", "Covariate; married/living with partner indicator"),
    ("total_calories", "Total energy intake", "Covariate; first-day dietary recall"),
    ("alcohol_drinks", "Alcohol intake", "Covariate"),
    ("is_ever_smoker", "Smoking history", "Covariate; ever-smoker indicator"),
    ("log_nlr", "log(NLR)", "Covariate; systemic inflammation marker"),
    ("sedentary_min", "Sedentary time", "Covariate"),
    ("is_diabetic", "Diabetes", "Covariate"),
    ("has_cvd", "Cardiovascular disease", "Covariate"),
    ("is_ad", "Antidepressant use", "Covariate/effect modifier"),
    ("is_statin", "Statin use", "Covariate"),
    ("has_liver_disease", "History of liver disease", "Covariate"),
    ("race", "Race/ethnicity", "Covariate/effect modifier"),
    ("WTMEC2YR", "MEC examination weight", "Survey weight; non-missing by analytic inclusion"),
]


DPQ_LABELS = {
    "DPQ010": "Little interest or pleasure",
    "DPQ020": "Feeling down, depressed, or hopeless",
    "DPQ030": "Sleep problems",
    "DPQ040": "Feeling tired or having little energy",
    "DPQ050": "Poor appetite or overeating",
    "DPQ060": "Feeling bad about yourself",
    "DPQ070": "Trouble concentrating",
    "DPQ080": "Moving/speaking slowly or being fidgety/restless",
    "DPQ090": "Thoughts of self-harm",
}


def pct(n_missing: int, denominator: int) -> str:
    return f"{100 * n_missing / denominator:.1f}%"


def build_missingness_dataframe() -> pd.DataFrame:
    pipeline = NHANESDepressionAnalysis(
        dataset_dir=str(ROOT / "data" / "raw"),
        n_jobs=1,
        imputer_max_iter=30,
    )
    pipeline.load_multicycle_data(cycles=("E", "F", "G", "H", "I", "J"))
    pipeline.preprocess()

    df = pipeline.df.dropna(subset=["WTMEC2YR"]).copy()
    denominator = len(df)
    if denominator != 36580:
        print(f"Warning: analytic denominator is {denominator}, expected 36580.")

    rows = []
    rows.append(
        {
            "Variable": "Analytic sample",
            "Role / definition": "Adults aged >=18 years with non-missing MEC examination weight",
            "Available, n": f"{denominator:,}",
            "Missing, n": "0",
            "Missing, %": "0.0%",
        }
    )

    dpq_cols = [f"DPQ{str(i).zfill(3)}" for i in range(10, 100, 10)]
    any_phq_missing = int(df[dpq_cols].isna().any(axis=1).sum())
    rows.append(
        {
            "Variable": "Any PHQ-9 item missing",
            "Role / definition": "Outcome construction; PHQ-9 total score calculated after item-level MICE",
            "Available, n": f"{denominator - any_phq_missing:,}",
            "Missing, n": f"{any_phq_missing:,}",
            "Missing, %": pct(any_phq_missing, denominator),
        }
    )

    for col in dpq_cols:
        missing = int(df[col].isna().sum())
        rows.append(
            {
                "Variable": col,
                "Role / definition": f"PHQ-9 item: {DPQ_LABELS[col]}",
                "Available, n": f"{denominator - missing:,}",
                "Missing, n": f"{missing:,}",
                "Missing, %": pct(missing, denominator),
            }
        )

    for col, label, role in VARIABLES:
        if col not in df.columns:
            missing = denominator
        else:
            missing = int(df[col].isna().sum())
        rows.append(
            {
                "Variable": label,
                "Role / definition": role,
                "Available, n": f"{denominator - missing:,}",
                "Missing, n": f"{missing:,}",
                "Missing, %": pct(missing, denominator),
            }
        )

    return pd.DataFrame(rows)


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    lines = [
        "| Variable | Role / definition | Available, n | Missing, n | Missing, % |",
        "|---|---|---:|---:|---:|",
    ]
    for _, row in df.iterrows():
        values = [
            str(row["Variable"]),
            str(row["Role / definition"]),
            str(row["Available, n"]),
            str(row["Missing, n"]),
            str(row["Missing, %"]),
        ]
        values = [v.replace("|", "\\|") for v in values]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines) + "\n"


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    apply_fonts_to_run(run, size_pt=9)
    run.italic = True


def build_docx(df: pd.DataFrame, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("S7 Table. Missingness of primary analysis variables before multiple imputation.")
    apply_fonts_to_run(run, bold=True, size_pt=13)

    add_note(
        doc,
        "Denominator is the analytic adult sample with non-missing MEC examination weights (N=36,580). "
        "Missingness is shown before MICE. Education missingness was retained as a separate category in the main model; PHQ-9 total score was calculated after item-level imputation.",
    )

    rows = [list(df.columns)] + df.astype(str).values.tolist()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = None
    apply_three_line_table_styling(table, rows)
    for row in table.rows:
        row.allow_break_across_pages = True

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    df = build_missingness_dataframe()
    md = dataframe_to_markdown(df)
    for path in [OUT_MD, PACKAGE_MD, TP_PACKAGE_MD]:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(md, encoding="utf-8")
    for path in [OUT_DOCX, PACKAGE_DOCX, TP_PACKAGE_DOCX]:
        build_docx(df, path)
    print(f"Saved {OUT_MD}")
    print(f"Saved {OUT_DOCX}")
    print(f"Saved {PACKAGE_MD}")
    print(f"Saved {PACKAGE_DOCX}")
    print(f"Saved {TP_PACKAGE_MD}")
    print(f"Saved {TP_PACKAGE_DOCX}")


if __name__ == "__main__":
    main()
