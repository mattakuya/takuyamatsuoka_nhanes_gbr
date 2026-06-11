import os
import glob
import re
import shutil
import math
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "results" / "manuscript_english.txt"
OUTPUT = ROOT / "submission_packages" / "translational_psychiatry" / "manuscript_tp.md"

def strip_markdown_italics(text: str) -> str:
    # Remove single asterisks around statistical variables (e.g. *N*, *β*, *p*, *df*, *χ*, *α*, *m*, *q*)
    return re.sub(r'\*([a-zA-Zβqαχ²\-]+)\*', r'\1', text)

def replace_thousands_commas_with_spaces(text: str) -> str:
    """Replace thousands separator commas with spaces in numbers like 36,580 or 92,795.85."""
    pattern = re.compile(r'\b(\d{1,3}),(\d{3})\b')
    while True:
        new_text = pattern.sub(r'\1 \2', text)
        if new_text == text:
            break
        text = new_text
    return text

def renumber_references_by_citation_order(text: str) -> str:
    """Renumber numeric references in order of first citation."""
    split_match = re.search(r"\nReferences\n", text)
    if not split_match:
        return text

    body = text[:split_match.end()]
    references_block = text[split_match.end():].strip()

    # Find if there are tables/sections after the references
    tables_match = re.search(r"\n####\s+Table", references_block)
    if tables_match:
        tables_part = references_block[tables_match.start():]
        refs_part = references_block[:tables_match.start()]
    else:
        tables_part = ""
        refs_part = references_block

    ref_entries = {}
    for line in refs_part.splitlines():
        match = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
        if match:
            ref_entries[int(match.group(1))] = match.group(2).strip()

    if not ref_entries:
        return text

    old_to_new = {}
    ordered_old_numbers = []

    def is_citation_token(token: str) -> bool:
        return bool(re.fullmatch(r"\s*\d+\s*", token))

    def replace_citation(match: re.Match) -> str:
        content = match.group(1)
        parts = [part.strip() for part in content.split(",")]
        if not parts or not all(is_citation_token(part) for part in parts):
            return match.group(0)

        new_numbers = []
        for part in parts:
            old_number = int(part)
            if old_number not in ref_entries:
                return match.group(0)
            if old_number not in old_to_new:
                old_to_new[old_number] = len(old_to_new) + 1
                ordered_old_numbers.append(old_number)
            new_numbers.append(str(old_to_new[old_number]))
        new_numbers = sorted(new_numbers, key=lambda value: int(value))
        return "[" + ", ".join(new_numbers) + "]"

    citation_pattern = re.compile(r"\[((?:\s*\d+\s*,)*\s*\d+\s*)\]")
    body = citation_pattern.sub(replace_citation, body)

    for old_number in sorted(ref_entries):
        if old_number not in old_to_new:
            old_to_new[old_number] = len(old_to_new) + 1
            ordered_old_numbers.append(old_number)

    renumbered_refs = [
        f"{old_to_new[old_number]}. {ref_entries[old_number]}"
        for old_number in ordered_old_numbers
    ]
    result = body.rstrip() + "\n" + "\n".join(renumbered_refs) + "\n"
    if tables_part:
        result += "\n" + tables_part.strip() + "\n"
    return result

def get_latest_results_file(results_dir=None):
    if results_dir is None:
        results_dir = str(ROOT / "results")
    pattern = os.path.join(results_dir, "nhanes_results_*.md")
    files = glob.glob(pattern)
    files = [f for f in files if "QUICK" not in f]
    if not files:
        raise FileNotFoundError(f"Results file matching 'nhanes_results_*.md' not found in {results_dir}")
    return max(files, key=os.path.getmtime)

def parse_markdown_table(file_path, header_title):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Locate the heading
    pattern = rf"##\s+{re.escape(header_title)}\s*\n\n_FDR family:.*_\n\n(\|[\s\S]*?)(?=\n\n##|\Z)"
    match = re.search(pattern, content)
    if not match:
        pattern_simple = rf"##\s+{re.escape(header_title)}\s*\n\n(\|[\s\S]*?)(?=\n\n##|\Z)"
        match = re.search(pattern_simple, content)
    
    if match:
        return match.group(1).strip()
    else:
        raise ValueError(f"Could not find table under heading: {header_title}")

def build_table_1(results_file):
    return parse_markdown_table(results_file, "DESCRIPTIVE TABLE 1 BY LOG_RATIO QUARTILE")

def build_table_2(results_file):
    main_table = parse_markdown_table(results_file, "MAIN ANALYSIS: PHQ-9 CONTINUOUS")
    logistic_table = parse_markdown_table(results_file, "SENSITIVITY: PHQ-9 ≥ 10 (Logistic, β on log-odds)")
    sensitivity_table = parse_markdown_table(results_file, "SENSITIVITY ANALYSIS")
    liver_safe_table = parse_markdown_table(results_file, "LIVER-SAFE (NO DILI) ANALYSIS")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Analysis Model / Target | N | Beta / OR [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    
    # 1. Primary WLS
    rows.append("| **Primary Analysis (WLS)** | | | | | |")
    main_row = extract_rows(main_table)[0].split('|')[1:-1]
    subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = [x.strip() for x in main_row]
    rows.append(f"| Full Sample (OLS PHQ-9) | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
    
    # 2. Logistic PHQ-9 >= 10
    rows.append("| **Sensitivity Analysis (Logistic)** | | | | | |")
    log_row = extract_rows(logistic_table)[0].split('|')[1:-1]
    subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = [x.strip() for x in log_row]
    or_val = math.exp(float(beta))
    or_lo = math.exp(float(ci_lo))
    or_hi = math.exp(float(ci_hi))
    rows.append(f"| Full Sample (Logistic PHQ-9 &ge; 10) | {n} | OR = {or_val:.4f} [{or_lo:.4f}, {or_hi:.4f}] | — | {p} | {q} |")
    
    # 3. Exclusion sensitivity
    rows.append("| **Sensitivity Analyses (Exclusion Models)** | | | | | |")
    for row_str in extract_rows(sensitivity_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    # 4. Liver-safe
    rows.append("| **Liver-Safe Analysis (No DILI)** | | | | | |")
    for row_str in extract_rows(liver_safe_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def build_table_3(results_file):
    bmi_table = parse_markdown_table(results_file, "BMI STRATIFIED ANALYSIS")
    sex_table = parse_markdown_table(results_file, "SEX-STRATIFIED ANALYSIS")
    age_table = parse_markdown_table(results_file, "AGE-STRATIFIED ANALYSIS")
    race_table = parse_markdown_table(results_file, "RACE-STRATIFIED ANALYSIS")
    alcohol_table = parse_markdown_table(results_file, "ALCOHOL-STRATIFIED ANALYSIS (NIAAA)")
    cotinine_table = parse_markdown_table(results_file, "COTININE-STRATIFIED ANALYSIS (Pirkle)")
    crp_table = parse_markdown_table(results_file, "CRP-STRATIFIED ANALYSIS (AHA, I/J only)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Subgroup / Stratification | N | Beta [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    
    def append_group(title, table_str, skip_full=True):
        rows.append(f"| **{title}** | | | | | |")
        for row_str in extract_rows(table_str):
            parts = [x.strip() for x in row_str.split('|')[1:-1]]
            subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
            if skip_full and "Full Sample" in subgroup:
                continue
            rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
            
    append_group("BMI Categories (kg/m²)", bmi_table)
    append_group("Sex", sex_table)
    append_group("Age Groups", age_table)
    append_group("Race / Ethnicity", race_table)
    append_group("Alcohol Consumption (NIAAA)", alcohol_table)
    append_group("Cotinine Level (Smoking Exposure)", cotinine_table)
    append_group("Systemic Inflammation (hs-CRP, I/J only)", crp_table)
    
    return '\n'.join([header, separator] + rows)

def build_table_4(results_file):
    obs_table = parse_markdown_table(results_file, "JOINT MODEL: log_ratio + OBS (per 1 SD each, SD-standardized)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Subgroup / Model | Term | N | Beta [95% CI] | SE | P-value | Attenuation (%) |"
    separator = "| :--- | :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    for row_str in extract_rows(obs_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        if len(parts) == 13:
            subgroup, model, term_part1, term_part2, n, beta, se, ci_lo, ci_hi, p, atten, q, q_fam = parts
            term = f"{term_part1} + {term_part2}"
        else:
            subgroup, model, term, n, beta, se, ci_lo, ci_hi, p, atten, q, q_fam = parts
        rows.append(f"| {subgroup} ({model}) | {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {atten} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_1(results_file):
    return parse_markdown_table(results_file, "VIF (Main Model Covariates)")

def build_supp_table_2(results_file):
    model_fit_table = parse_markdown_table(results_file, "MODEL FIT COMPARISON: RATIO vs GGT vs GGT+BIL")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Model | N | Parameters | Pseudo-AIC | Pseudo-BIC | Adjusted R² | Weighted RMSE |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    for row_str in extract_rows(model_fit_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        model, n, m, aic_mean, aic_sd, bic_mean, bic_sd, r2_mean, r2_sd, rmse_mean, rmse_sd = parts
        rows.append(f"| {model} | {n} | {m} | {float(aic_mean):.2f} | {float(bic_mean):.2f} | {float(r2_mean):.4f} | {float(rmse_mean):.4f} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_3(results_file):
    return parse_markdown_table(results_file, "CCA vs MICE: GGT BIAS ASSESSMENT (SD-standardized)")

def build_supp_table_4(results_file):
    spec_table = parse_markdown_table(results_file, "OXIDATIVE COMPARISON: GGT/Bil vs CDAI vs OBS (SD-STANDARDIZED)")
    excl_ad_table = parse_markdown_table(results_file, "OXIDATIVE COMPARISON EXCL. AD USERS (SD-STANDARDIZED)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Analysis Cohort / Indicator | N | Beta [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    rows.append("| **Full Cohort (OBS-Restricted Sample)** | | | | | |")
    for row_str in extract_rows(spec_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    rows.append("| **Antidepressant Non-Users Cohort** | | | | | |")
    for row_str in extract_rows(excl_ad_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_5(results_file):
    ad_table = parse_markdown_table(results_file, "AD-INTERACTION ANALYSIS [log_ratio]")
    ad_normal_table = parse_markdown_table(results_file, "AD-INTERACTION ANALYSIS [log_ratio] (Normal BMI)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Sample / Term | N | Beta [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    rows.append("| **Full Cohort Interaction Model** | | | | | |")
    for row_str in extract_rows(ad_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        term, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    rows.append("| **Normal BMI (18.5–24.9 kg/m²) Cohort** | | | | | |")
    for row_str in extract_rows(ad_normal_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        term, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_6(race_results_file):
    race_interaction = parse_markdown_table(race_results_file, "RACE/ETHNICITY INTERACTION ANALYSIS [log_ratio]")
    global_test = parse_markdown_table(race_results_file, "GLOBAL RACE/ETHNICITY INTERACTION TEST [log_ratio]")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Term / Test | N | Beta / Chi-square [95% CI / df] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    rows.append("| **Interaction Model (NHWhite as Reference)** | | | | | |")
    for row_str in extract_rows(race_interaction):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        term, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    rows.append("| **Global Interaction Wald Test** | | | | | |")
    for row_str in extract_rows(global_test):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        test, n, chi2, df, p, terms, q, q_fam = parts
        rows.append(f"| {test} | {n} | Chi² = {float(chi2):.4f} [df = {df}] | — | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def apply_fonts_to_run(run, bold=False, size_pt=11):
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Helvetica')
    rFonts.set(qn('w:hAnsi'), 'Helvetica')
    rPr.append(rFonts)
    
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True

def apply_three_line_table_styling(table, rows_data):
    """
    Applies professional booktabs (three-line table) styling to a python-docx table.
    - No vertical lines.
    - Thick top rule (1.5 pt = 12 sz) and bottom rule (1.5 pt = 12 sz).
    - Thin header rule below row 0 (0.75 pt = 6 sz).
    - Cell padding (5 pt top/bottom, 7.5 pt left/right).
    - Merged cells and light gray shading (#F2F2F2) for group header rows.
    """
    tblPr = table._tbl.tblPr
    
    # 1. Clear standard borders by adding custom w:tblBorders
    tblBorders = OxmlElement('w:tblBorders')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '111111')
    tblBorders.append(top)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'none')
    tblBorders.append(bottom)
    
    for border_name in ('left', 'right', 'insideV', 'insideH'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
        
    tblPr.append(tblBorders)
    
    # 2. Set table-level cell padding (dxa unit)
    tblCellMar = OxmlElement('w:tblCellMar')
    for margin_name, val in [('top', 100), ('bottom', 100), ('left', 150), ('right', 150)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)
    tblPr.append(tblCellMar)
    
    # 3. Render and style cells
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        
        # Check if this is a group header row (all bold or empty other cells)
        is_group_header = False
        first_val = row_data[0].strip()
        if r_idx > 0 and (first_val.startswith('**') and first_val.endswith('**')):
            is_group_header = True
        elif r_idx > 0 and all(cell.strip() == '' for cell in row_data[1:]):
            is_group_header = True
            
        if is_group_header:
            # Merge all cells in this row
            merged_cell = row.cells[0]
            for c_idx in range(1, len(row.cells)):
                merged_cell = merged_cell.merge(row.cells[c_idx])
            
            # Style the merged group header cell
            merged_cell.text = ""
            p = merged_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(first_val.replace('**', ''))
            apply_fonts_to_run(run, bold=True, size_pt=10)
            
            # Subtle background color (#F2F2F2)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'F2F2F2')
            merged_cell._tc.get_or_add_tcPr().append(shading)
            continue
            
        for c_idx, cell_value in enumerate(row_data):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                # Alignments
                if c_idx == 0:
                    p.alignment = 0  # Left
                else:
                    p.alignment = 1  # Center
                
                is_header = (r_idx == 0)
                clean_val = cell_value.replace("**", "").replace("&ge;", "≥").replace("&le;", "≤")
                
                run = p.add_run(clean_val)
                apply_fonts_to_run(run, bold=is_header, size_pt=10)
                
                # Add thin bottom border to header row (row 0)
                if r_idx == 0:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')  # 0.75 pt
                    bottom.set(qn('w:space'), '0')
                    bottom.set(qn('w:color'), '111111')
                    tcBorders.append(bottom)
                    tcPr.append(tcBorders)

                # Add thick bottom border to last row (bottom rule)
                if r_idx == len(rows_data) - 1:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '12')  # 1.5 pt
                    bottom.set(qn('w:space'), '0')
                    bottom.set(qn('w:color'), '111111')
                    tcBorders.append(bottom)
                    tcPr.append(tcBorders)

    # 4. Prevent rows from splitting across pages
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
        
    in_table = False
    table_lines = []
    
    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        
        rows_data = []
        for line in table_lines:
            if not line.strip() or re.match(r'^\s*\|\s*[:\-|\s]+$', line):
                continue
            parts = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            rows_data.append(parts)
            
        if not rows_data:
            table_lines = []
            return
            
        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = None
        apply_three_line_table_styling(table, rows_data)
        
        doc.add_paragraph()
        table_lines = []
        
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.strip().startswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False
            
        if not line.strip():
            i += 1
            continue
            
        h1_match = re.match(r'^#\s+(.*)', line)
        if h1_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(h1_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=16)
            i += 1
            continue
            
        h2_match = re.match(r'^##\s+(.*)', line)
        if h2_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h2_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=13)
            i += 1
            continue
            
        h3_match = re.match(r'^###\s+(.*)', line)
        if h3_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h3_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=11)
            i += 1
            continue
            
        h4_match = re.match(r'^####\s+(.*)', line)
        if h4_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h4_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=10.5)
            i += 1
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                bold_text = part[2:-2]
                run = p.add_run(bold_text)
                apply_fonts_to_run(run, bold=True, size_pt=11)
            else:
                run = p.add_run(part)
                apply_fonts_to_run(run, bold=False, size_pt=11)
                
        i += 1
        
    if in_table:
        flush_table()
        
    doc.save(docx_path)
    print(f"Successfully saved Word file to: {docx_path}")

def replace_once(text: str, old: str, new: str) -> str:
    if old not in text:
        raise ValueError(f"Target block not found:\n{old[:300]}")
    return text.replace(old, new, 1)

def main() -> None:
    results_file = get_latest_results_file()
    race_results_file = results_file
    print(f"Parsing raw results from: {results_file}")

    # Build main tables
    table1 = build_table_1(results_file)
    table2 = build_table_2(results_file)
    table3 = build_table_3(results_file)
    table4 = build_table_4(results_file)

    # Build supplementary tables
    supp_table1 = build_supp_table_1(results_file)
    supp_table2 = build_supp_table_2(results_file)
    supp_table3 = build_supp_table_3(results_file)
    supp_table4 = build_supp_table_4(results_file)
    supp_table5 = build_supp_table_5(results_file)
    supp_table6 = build_supp_table_6(race_results_file)
    missingness_path = ROOT / "results" / "supp_table7_missingness.md"
    supp_table7 = (
        missingness_path.read_text(encoding="utf-8").strip()
        if missingness_path.exists()
        else "_Run generate_missingness_table.py to create this table._"
    )

    # Read original draft template
    text = SOURCE.read_text(encoding="utf-8")

    # Perform manuscript edits for Translational Psychiatry
    text = replace_once(
        text,
        "Race- and ethnicity-specific association of the gamma-glutamyl transferase-to-bilirubin ratio with depressive symptoms in US adults: a NHANES 2007–2018 study",
        "Race- and ethnicity-specific association of the gamma-glutamyl transferase-to-bilirubin ratio with depressive symptoms in US adults: a NHANES 2007–2018 study\n\nRunning title: Redox balance and depressive symptoms",
    )
    text = replace_once(
        text,
        "¹ Tohoku University Hospital, Sendai, Miyagi, Japan",
        "¹ Tohoku University Hospital, Sendai, Japan",
    )

    # Normalize section headings to Nature style
    text = text.replace("Materials and methods", "Methods", 1)
    text = text.replace("Supporting information", "Supplementary Information", 1)

    # Replace table placeholders with the TP-style [Insert Table X here]
    text = text.replace("{table1}", "[Insert Table 1 here]")
    text = text.replace("{table2}", "[Insert Table 2 here]")
    text = text.replace("{table3}", "[Insert Table 3 here]")
    text = text.replace("{table4}", "[Insert Table 4 here]")

    # Strip markdown-like asterisks around statistical terms
    text = strip_markdown_italics(text)

    # Normalize figure/table citation labels to Nature Portfolio / Translational Psychiatry style
    text = re.sub(r'\bS(\d+)\s+Fig\b', r'Supplementary Fig. S\1', text)
    text = re.sub(r'\bS(\d+)\s+Table\b', r'Supplementary Table S\1', text)
    text = re.sub(r'\bS(\d+)\s+Checklist\b', r'Supplementary Checklist S\1', text)
    text = re.sub(r'\bFig\s+(\d+)\b', r'Fig. \1', text)

    text = renumber_references_by_citation_order(text)
    
    # Replace thousands separator commas with spaces to conform to TP House Style
    text = replace_thousands_commas_with_spaces(text)

    # Save Markdown Main Manuscript
    OUTPUT.write_text(text, encoding="utf-8")
    print(f"Successfully compiled and saved English manuscript to: {OUTPUT}")

    # Generate Supplementary Material Markdown
    supp_text = f"""# Race- and ethnicity-specific association of the gamma-glutamyl transferase-to-bilirubin ratio with depressive symptoms in US adults: a NHANES 2007–2018 study
## Supplementary Materials

**Takuya Matsuoka¹***
¹ Tohoku University Hospital, Sendai, Japan
* Corresponding author (email: takuya.matsuoka.c6@tohoku.ac.jp)

---

### Supplementary Tables

#### S1 Table. Multicollinearity diagnostics for main model covariates
{supp_table1}

#### S2 Table. Model fit comparison of the primary exposure variables
{supp_table2}

#### S3 Table. Complete case analysis versus multiple imputation by chained equations (MICE)
{supp_table3}

#### S4 Table. Associations with established oxidative stress indicators and joint models
{supp_table4}

#### S5 Table. Antidepressant interaction analyses
{supp_table5}

#### S6 Table. Race/ethnicity interaction analysis
{supp_table6}

#### S7 Table. Missingness of primary analysis variables before multiple imputation
{supp_table7}

---
"""
    # Strip markdown-like asterisks from supplementary materials
    supp_text = strip_markdown_italics(supp_text)
    
    # Replace thousands separator commas with spaces in supplementary materials
    supp_text = replace_thousands_commas_with_spaces(supp_text)

    # Normalize supplementary material headings to Nature style
    supp_text = re.sub(r'\bS(\d+)\s+Fig\b', r'Supplementary Fig. S\1', supp_text)
    supp_text = re.sub(r'\bS(\d+)\s+Table\b', r'Supplementary Table S\1', supp_text)
    supp_text = re.sub(r'\bS(\d+)\s+Checklist\b', r'Supplementary Checklist S\1', supp_text)

    out_dir = ROOT / "submission_packages" / "translational_psychiatry"
    out_dir.mkdir(parents=True, exist_ok=True)
    supp_output_path = out_dir / "supplementary_material_tp.md"
    supp_output_path.write_text(supp_text, encoding="utf-8")
    print(f"Successfully compiled and saved English supplementary material to: {supp_output_path}")

    # Convert to Word documents
    convert_md_to_docx(str(OUTPUT), str(out_dir / "manuscript_tp.docx"))
    convert_md_to_docx(str(supp_output_path), str(out_dir / "supplementary_material_tp.docx"))
    print("All English Translational Psychiatry submissions compiled successfully!")

if __name__ == "__main__":
    main()
