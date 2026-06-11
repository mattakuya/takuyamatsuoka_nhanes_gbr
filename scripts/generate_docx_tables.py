import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Add project root to path to ensure imports work
sys.path.append(str(Path(__file__).resolve().parent.parent))
sys.path.append(str(Path(__file__).resolve().parent))

from generate_tp_manuscript import (
    get_latest_results_file,
    parse_markdown_table,
    build_table_1,
    build_table_2,
    build_table_3,
    build_table_4,
    apply_fonts_to_run,
    apply_three_line_table_styling
)

def set_column_widths(table, widths_inches):
    """
    Sets column widths cell-by-cell in inches.
    Only sets cell width for rows that do not contain merged cells.
    """
    for row in table.rows:
        if len(row.cells) == len(widths_inches):
            for c_idx, width in enumerate(widths_inches):
                row.cells[c_idx].width = Inches(width)

def main():
    try:
        results_file = get_latest_results_file()
        print(f"Reading results from: {results_file}")
        
        output_dir = Path("submission_packages/translational_psychiatry")
        output_dir.mkdir(parents=True, exist_ok=True)
        docx_path = output_dir / "tables_tp.docx"
        
        print("Creating beautiful standalone tables document...")
        doc = Document()
        
        # Configure landscape orientation for all sections to accommodate wide clinical tables
        for section in doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap page dimensions for landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Helper to parse markdown table into docx
        def add_beautiful_table(title, md_table_str, col_widths=None, note=None):
            # Title
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(12)
            p_title.paragraph_format.space_after = Pt(8)
            run = p_title.add_run(title)
            apply_fonts_to_run(run, bold=True, size_pt=12)
            
            # Split markdown lines
            lines = md_table_str.strip().split('\n')
            
            # Filter out separator lines
            rows_data = []
            for line in lines:
                if not line.strip() or re.match(r'^\s*\|\s*[:\-|\s]+$', line):
                    continue
                parts = [cell.strip() for cell in line.strip().split('|')[1:-1]]
                rows_data.append(parts)
                
            if not rows_data:
                print(f"Warning: No rows found for table: {title}")
                return
                
            num_cols = max(len(row) for row in rows_data)
            num_rows = len(rows_data)
            
            # Create table
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = None
            
            # Apply standard premium three-line styling
            apply_three_line_table_styling(table, rows_data)
            
            # Apply custom column widths if provided
            if col_widths:
                set_column_widths(table, col_widths)

            if note:
                p_note = doc.add_paragraph()
                p_note.paragraph_format.space_before = Pt(6)
                p_note.paragraph_format.space_after = Pt(6)
                run_note = p_note.add_run(note)
                apply_fonts_to_run(run_note, size_pt=9)
                
            # Add a breathing space paragraph
            doc.add_paragraph()
            
        # 1. Table 1
        print("Compiling Table 1...")
        table1_md = build_table_1(results_file)
        add_beautiful_table(
            "Table 1. Participant characteristics by quartiles of log(GBR).",
            table1_md,
            col_widths=[3.5, 1.1, 1.1, 1.1, 1.1, 1.1],
            note="Values are unweighted counts, means (SD), or weighted percentages, as indicated. GBR, gamma-glutamyl transferase-to-total bilirubin ratio."
        )
        doc.add_page_break()
        
        # 2. Table 2
        print("Compiling Table 2...")
        table2_md = build_table_2(results_file)
        add_beautiful_table(
            "Table 2. Primary and sensitivity analyses of the association between log(GBR) and depressive symptoms.",
            table2_md,
            col_widths=[3.5, 0.9, 2.4, 0.8, 0.7, 0.7],
            note="Estimates are from fully adjusted weighted models unless otherwise indicated. The logistic model reports odds ratios for PHQ-9 >= 10."
        )
        doc.add_page_break()
        
        # 3. Table 3
        print("Compiling Table 3...")
        table3_md = build_table_3(results_file)
        add_beautiful_table(
            "Table 3. Subgroup analyses of the association between GBR and depressive symptoms.",
            table3_md,
            col_widths=[3.2, 0.9, 2.4, 0.8, 0.8, 0.9]
        )
        doc.add_page_break()
        
        # 4. Table 4
        print("Compiling Table 4...")
        table4_md = build_table_4(results_file)
        add_beautiful_table(
            "Table 4. Comparison of GBR with established oxidative stress indicators and joint models.",
            table4_md,
            col_widths=[2.6, 2.0, 0.7, 1.7, 0.6, 0.7, 0.7]
        )
        
        # Save Standalone Document
        doc.save(str(docx_path))
        print(f"Successfully compiled and saved separate tables file to: {docx_path}")
        
    except Exception as e:
        print(f"Error compiling standalone docx tables: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
