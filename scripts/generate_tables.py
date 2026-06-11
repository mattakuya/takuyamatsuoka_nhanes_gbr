import os
import re
import sys
import pandas as pd
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
import matplotlib.pyplot as plt

from utils import (
    get_latest_results_file,
    parse_markdown_table,
    build_table_1,
    build_table_2,
    build_table_3,
    build_table_4,
    build_table_5,
    apply_fonts_to_run,
    apply_three_line_table_styling,
    markdown_to_dataframe
)

OUTPUT_DIR = "results"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ==============================================================================
# 1. Standalone tables.docx Generation
# ==============================================================================
def set_column_widths(table, widths_inches):
    for row in table.rows:
        if len(row.cells) == len(widths_inches):
            for c_idx, width in enumerate(widths_inches):
                row.cells[c_idx].width = Inches(width)

def generate_docx_tables(results_file):
    print("Compiling standalone Word tables (results/tables.docx)...")
    doc = Document()
    
    # Configure landscape orientation for wide tables
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    def add_beautiful_table(title, md_table_str, col_widths=None, note=None):
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(8)
        run = p_title.add_run(title)
        apply_fonts_to_run(run, bold=True, size_pt=12)
        
        lines = md_table_str.strip().split('\n')
        rows_data = []
        for line in lines:
            if not line.strip() or re.match(r'^\s*\|\s*[:\-|\s]+$', line):
                continue
            parts = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            rows_data.append(parts)
            
        if not rows_data:
            return
            
        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = None
        apply_three_line_table_styling(table, rows_data)
        
        if col_widths:
            set_column_widths(table, col_widths)

        if note:
            p_note = doc.add_paragraph()
            p_note.paragraph_format.space_before = Pt(6)
            p_note.paragraph_format.space_after = Pt(6)
            run_note = p_note.add_run(note)
            apply_fonts_to_run(run_note, size_pt=9)
            
        doc.add_paragraph()
        
    # Table 1
    table1_md = build_table_1(results_file)
    add_beautiful_table(
        "Table 1. Participant characteristics by quartiles of log(GBR).",
        table1_md,
        col_widths=[3.5, 1.1, 1.1, 1.1, 1.1, 1.1],
        note="Values are unweighted counts, means (SD), or weighted percentages, as indicated. GBR, gamma-glutamyl transferase-to-total bilirubin ratio."
    )
    doc.add_page_break()
    
    # Table 2
    table2_md = build_table_2(results_file)
    add_beautiful_table(
        "Table 2. Primary and sensitivity analyses of the association between log(GBR) and depressive symptoms.",
        table2_md,
        col_widths=[3.5, 0.9, 2.4, 0.8, 0.7, 0.7],
        note="Estimates are from fully adjusted weighted models unless otherwise indicated. The logistic model reports odds ratios for PHQ-9 >= 10."
    )
    doc.add_page_break()
    
    # Table 3
    table3_md = build_table_3(results_file)
    add_beautiful_table(
        "Table 3. Subgroup analyses of the association between GBR and depressive symptoms.",
        table3_md,
        col_widths=[3.2, 0.9, 2.4, 0.8, 0.8, 0.9]
    )
    doc.add_page_break()
    
    # Table 4
    table4_md = build_table_4(results_file)
    add_beautiful_table(
        "Table 4. Comparison of GBR with established oxidative stress indicators and joint models.",
        table4_md,
        col_widths=[2.6, 2.0, 0.7, 1.7, 0.6, 0.7, 0.7]
    )
    
    docx_path = os.path.join(OUTPUT_DIR, "tables.docx")
    doc.save(docx_path)
    print(f"  Word tables file successfully saved to: {docx_path}")

# ==============================================================================
# 2. Table Image Rendering (PNG)
# ==============================================================================
def render_table_image(df, title, output_path, figsize=(10, 6), col_widths=None):
    fig, ax = plt.subplots(figsize=figsize)
    ax.axis('off')
    
    plt.rcParams['font.sans-serif'] = ['Helvetica', 'Arial', 'Hiragino Sans', 'AppleGothic', 'sans-serif']
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['font.size'] = 10
    
    ax.set_title(title, fontsize=11, weight='bold', pad=20, loc='left', color='#2c3e50')
    
    tbl = ax.table(
        cellText=df.values,
        colLabels=df.columns,
        cellLoc='left',
        loc='center',
        colWidths=col_widths
    )
    
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(8.5)
    tbl.scale(1.0, 1.6)
    
    n_rows = len(df)
    border_color_main = '#2c3e50'
    border_color_light = '#dcdde1'
    bg_section = '#f8f9fa'
    
    for (row_idx, col_idx), cell in tbl.get_celld().items():
        if row_idx == 0:
            cell.visible_edges = 'TB'
        elif row_idx == n_rows:
            cell.visible_edges = 'B'
        else:
            cell.visible_edges = ''
        
        is_section = False
        if row_idx > 0:
            row_values = df.iloc[row_idx - 1]
            first_val = str(row_values.iloc[0]).strip()
            if first_val.startswith('**') and first_val.endswith('**'):
                is_section = True
            elif col_idx == 0 and all(str(x).strip() == '' for x in row_values.iloc[1:]):
                is_section = True
        
        if row_idx == 0:
            cell.set_facecolor('#ffffff')
            cell.set_text_props(weight='bold', color='#111111')
            cell.set_linewidth(1.5)
            cell.set_edgecolor(border_color_main)
            cell.set_height(0.06)
        elif is_section:
            cell.set_facecolor(bg_section)
            cell.set_text_props(weight='bold', color='#2c3e50')
            cell.set_linewidth(0.5)
            cell.set_edgecolor(border_color_light)
            clean_text = cell.get_text().get_text().replace('**', '')
            cell.get_text().set_text(clean_text)
        else:
            cell.set_facecolor('#ffffff')
            cell.set_text_props(color='#2c3e50')
            cell.set_linewidth(0.5)
            
            if row_idx == n_rows:
                cell.set_edgecolor(border_color_main)
                cell.set_linewidth(1.5)
            else:
                cell.set_edgecolor(border_color_light)
                
            if col_idx > 0:
                cell.set_text_props(horizontalalignment='center')
                
            text_obj = cell.get_text()
            orig_text = text_obj.get_text()
            clean_text = orig_text.replace('&ge;', '≥').replace('&le;', '≤')
            text_obj.set_text(clean_text)
            
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"  Table image saved to: {output_path}")

def generate_table_images(results_file):
    print("Generating Table Images (PNG)...")
    
    # 1. Table 1
    print("  Rendering Table 1 Image...")
    table1_md = build_table_1(results_file)
    df1 = markdown_to_dataframe(table1_md)
    render_table_image(
        df1, "Table 1. Participant Characteristics by log(GBR) Quartile (NHANES 2007-2018)",
        os.path.join(OUTPUT_DIR, "table1_characteristics.png"),
        figsize=(11, 8.5), col_widths=[0.35, 0.13, 0.13, 0.13, 0.13, 0.13]
    )
    
    # 2. Table 2
    print("  Rendering Table 2 Image...")
    table2_md = build_table_2(results_file)
    df2 = markdown_to_dataframe(table2_md)
    render_table_image(
        df2, "Table 2. Primary and Sensitivity Analyses of log(GBR) in the Full Sample",
        os.path.join(OUTPUT_DIR, "table2_association.png"),
        figsize=(12, 6.5), col_widths=[0.36, 0.08, 0.26, 0.08, 0.11, 0.11]
    )
    
    # 3. Table 3
    print("  Rendering Table 3 Image...")
    table3_md = build_table_3(results_file)
    df3 = markdown_to_dataframe(table3_md)
    render_table_image(
        df3, "Table 3. Subgroup Stratified Analyses of GBR and PHQ-9",
        os.path.join(OUTPUT_DIR, "table3_subgroups.png"),
        figsize=(13, 11), col_widths=[0.32, 0.08, 0.28, 0.08, 0.12, 0.12]
    )
    
    # 4. Table 4
    print("  Rendering Table 4 Image...")
    table4_md = build_table_4(results_file)
    df4 = markdown_to_dataframe(table4_md)
    render_table_image(
        df4, "Table 4. Specificity Analyses of GBR and Other Liver Enzymes (Standardized per 1 SD)",
        os.path.join(OUTPUT_DIR, "table4_specificity.png"),
        figsize=(12, 5.5), col_widths=[0.32, 0.08, 0.28, 0.08, 0.12, 0.12]
    )
    
    # 5. Table 5
    print("  Rendering Table 5 Image...")
    table5_md = build_table_5(results_file)
    df5 = markdown_to_dataframe(table5_md)
    render_table_image(
        df5, "Table 5. Independent and Additive Association of GBR and OBS (Standardized per 1 SD)",
        os.path.join(OUTPUT_DIR, "table5_gbr_obs.png"),
        figsize=(13, 5.5), col_widths=[0.32, 0.24, 0.08, 0.20, 0.06, 0.06, 0.04]
    )
    
    # 6. Supplementary Table 3 (hsCRP Sensitivity)
    print("  Rendering Supplementary Table 3 Image...")
    from utils import parse_markdown_table
    try:
        table_hscrp_md = parse_markdown_table(results_file, "I/J HS-CRP SENSITIVITY")
        df_hscrp = markdown_to_dataframe(table_hscrp_md)
        render_table_image(
            df_hscrp, "Supplementary Table 3. hs-CRP Sensitivity Analysis in cycles I/J",
            os.path.join(OUTPUT_DIR, "supp_table3_hscrp.png"),
            figsize=(11, 4.5), col_widths=[0.35, 0.12, 0.26, 0.12, 0.15]
        )
    except Exception as e:
        print(f"    Skipping hsCRP image: {e}")
        
    # 7. Supplementary Table 4 (CCA vs MICE)
    print("  Rendering Supplementary Table 4 Image...")
    try:
        table_ccamice_md = parse_markdown_table(results_file, "CCA vs MICE: GGT BIAS ASSESSMENT (SD-standardized)")
        df_ccamice = markdown_to_dataframe(table_ccamice_md)
        render_table_image(
            df_ccamice, "Supplementary Table 4. Complete Case Analysis (CCA) vs Multiple Imputation (MICE) Bias Assessment",
            os.path.join(OUTPUT_DIR, "supp_table4_ccamice.png"),
            figsize=(11, 4.5), col_widths=[0.25, 0.15, 0.12, 0.26, 0.12, 0.10]
        )
    except Exception as e:
        print(f"    Skipping CCA vs MICE image: {e}")

# ==============================================================================
# 3. Missingness Table Builder (evaluated from raw data if available)
# ==============================================================================
VARIABLES = [
    ("log_ratio", "log(GBR)", "Primary exposure; log(GGT/total bilirubin)"),
    ("LBXSGTSI", "Serum GGT", "Exposure component"),
    ("LBXSTB", "Total bilirubin", "Exposure component"),
    ("RIDAGEYR", "Age", "Covariate"),
    ("is_female", "Sex", "Covariate; female indicator"),
    ("BMXBMI", "BMI", "Covariate"),
    ("pir", "Poverty-to-income ratio", "Covariate"),
    ("education_cat", "Education category", "Covariate; missing category retained in models"),
    ("is_married", "Marital status", "Covariate; married/living with partner indicator"),
    ("total_calories", "Total energy intake", "Covariate; first-day dietary recall"),
    ("alcohol_drinks", "Alcohol intake", "Covariate"),
    ("is_ever_smoker", "Smoking history", "Covariate; ever-smoker indicator"),
    ("log_nlr", "log(NLR)", "Covariate; systemic inflammation marker"),
    ("sedentary_min", "Sedentary time", "Covariate"),
    ("is_diabetic", "Diabetes", "Covariate"),
    ("has_cvd", "Cardiovascular disease", "Covariate"),
    ("is_ad", "Antidepressant use", "Covariate/effect modifier"),
    ("is_statin", "Statin use", "Covariate"),
    ("has_liver_disease", "History of liver disease", "Covariate"),
    ("race", "Race/ethnicity", "Covariate/effect modifier"),
    ("WTMEC2YR", "MEC examination weight", "Survey weight; non-missing by analytic inclusion"),
]

DPQ_LABELS = {
    "DPQ010": "Little interest or pleasure",
    "DPQ020": "Feeling down, depressed, or hopeless",
    "DPQ030": "Sleep problems",
    "DPQ040": "Feeling tired or having little energy",
    "DPQ050": "Poor appetite or overeating",
    "DPQ060": "Feeling bad about yourself",
    "DPQ070": "Trouble concentrating",
    "DPQ080": "Moving/speaking slowly or being fidgety/restless",
    "DPQ090": "Thoughts of self-harm",
}

def pct(n_missing: int, denominator: int) -> str:
    return f"{100 * n_missing / denominator:.1f}%"

def build_missingness_dataframe(raw_data_dir) -> pd.DataFrame:
    from analysis_pipeline import NHANESDepressionAnalysis
    pipeline = NHANESDepressionAnalysis(
        dataset_dir=raw_data_dir,
        n_jobs=1,
        imputer_max_iter=30,
    )
    pipeline.load_multicycle_data(cycles=("E", "F", "G", "H", "I", "J"))
    pipeline.preprocess()

    df = pipeline.df.dropna(subset=["WTMEC2YR"]).copy()
    denominator = len(df)

    rows = []
    rows.append({
        "Variable": "Analytic sample",
        "Role / definition": "Adults aged >=18 years with non-missing MEC examination weight",
        "Available, n": f"{denominator:,}",
        "Missing, n": "0",
        "Missing, %": "0.0%",
    })

    dpq_cols = [f"DPQ{str(i).zfill(3)}" for i in range(10, 100, 10)]
    any_phq_missing = int(df[dpq_cols].isna().any(axis=1).sum())
    rows.append({
        "Variable": "Any PHQ-9 item missing",
        "Role / definition": "Outcome construction; PHQ-9 total score calculated after item-level MICE",
        "Available, n": f"{denominator - any_phq_missing:,}",
        "Missing, n": f"{any_phq_missing:,}",
        "Missing, %": pct(any_phq_missing, denominator),
    })

    for col in dpq_cols:
        missing = int(df[col].isna().sum())
        rows.append({
            "Variable": col,
            "Role / definition": f"PHQ-9 item: {DPQ_LABELS[col]}",
            "Available, n": f"{denominator - missing:,}",
            "Missing, n": f"{missing:,}",
            "Missing, %": pct(missing, denominator),
        })

    for col, label, role in VARIABLES:
        if col not in df.columns:
            missing = denominator
        else:
            missing = int(df[col].isna().sum())
        rows.append({
            "Variable": label,
            "Role / definition": role,
            "Available, n": f"{denominator - missing:,}",
            "Missing, n": f"{missing:,}",
            "Missing, %": pct(missing, denominator),
        })

    return pd.DataFrame(rows)

def dataframe_to_markdown(df: pd.DataFrame) -> str:
    lines = [
        "| Variable | Role / definition | Available, n | Missing, n | Missing, % |",
        "|---|---|---:|---:|---:|",
    ]
    for _, row in df.iterrows():
        values = [
            str(row["Variable"]),
            str(row["Role / definition"]),
            str(row["Available, n"]),
            str(row["Missing, n"]),
            str(row["Missing, %"]),
        ]
        values = [v.replace("|", "\\|") for v in values]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines) + "\n"

def build_missingness_docx(df: pd.DataFrame, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Supplementary Table S7. Missingness of primary analysis variables before multiple imputation.")
    apply_fonts_to_run(run, bold=True, size_pt=13)

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_after = Pt(6)
    run_note = p_note.add_run(
        "Denominator is the analytic adult sample with non-missing MEC examination weights (N=36,580). "
        "Missingness is shown before MICE. Education missingness was retained as a separate category in the main model; PHQ-9 total score was calculated after item-level imputation."
    )
    apply_fonts_to_run(run_note, size_pt=9)
    run_note.italic = True

    rows = [list(df.columns)] + df.astype(str).values.tolist()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = None
    apply_three_line_table_styling(table, rows)
    for row in table.rows:
        row.allow_break_across_pages = True

    doc.save(path)

def generate_missingness_table():
    print("Evaluating pre-imputation missingness table...")
    try:
        raw_data_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "raw")
        if not os.path.exists(raw_data_dir):
            raw_data_dir = os.path.join("data", "raw")
            
        if not os.path.exists(raw_data_dir):
            raise FileNotFoundError(f"Raw data directory not found at {raw_data_dir}. Skipping.")

        df = build_missingness_dataframe(raw_data_dir)
        md = dataframe_to_markdown(df)
        
        md_path = Path(OUTPUT_DIR) / "supp_table7_missingness.md"
        docx_path = Path(OUTPUT_DIR) / "supp_table7_missingness.docx"
        
        md_path.write_text(md, encoding="utf-8")
        build_missingness_docx(df, docx_path)
        print(f"  Successfully saved missingness table to: {md_path} and {docx_path}")
    except Exception as e:
        print(f"  Skipping missingness table (raw CDC data or pipeline required): {e}")

# ==============================================================================
# MAIN RUNNER
# ==============================================================================
def main():
    print("=== STARTING TABLE GENERATION PIPELINE ===")
    try:
        results_file = get_latest_results_file(OUTPUT_DIR)
        print(f"Using results file: {results_file}")
    except FileNotFoundError as e:
        print(f"Error: {e}")
        print("Please run scripts/analysis_pipeline.py first to produce nhanes_results_*.md.")
        sys.exit(1)
        
    generate_docx_tables(results_file)
    generate_table_images(results_file)
    generate_missingness_table()
    print("=== TABLE GENERATION PIPELINE COMPLETE ===")

if __name__ == "__main__":
    main()
