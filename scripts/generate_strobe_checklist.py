from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "submission_packages" / "medrxiv" / "strobe_checklist_cross_sectional.docx"
TP_OUT = ROOT / "submission_packages" / "translational_psychiatry" / "strobe_checklist_cross_sectional.docx"
MD_OUT = ROOT / "submission_packages" / "medrxiv" / "strobe_checklist_cross_sectional.md"
TP_MD_OUT = ROOT / "submission_packages" / "translational_psychiatry" / "strobe_checklist_cross_sectional.md"


CHECKLIST_ROWS = [
    ("Title and abstract", "1a", "Indicate the study design with a commonly used term in the title or abstract.", "Title; Abstract", "The title states that the study is cross-sectional; the abstract identifies NHANES 2007-2018 cross-sectional data."),
    ("Title and abstract", "1b", "Provide an informative and balanced summary of what was done and what was found.", "Abstract", "The abstract summarizes the background, data source, participants, exposure, outcome, analytic approach, main estimates, race/ethnicity findings, and cautious interpretation."),
    ("Introduction", "2", "Explain the scientific background and rationale for the investigation.", "Introduction", "The introduction describes oxidative stress in depression, limitations of specialized oxidative stress assays, existing OBS/CDAI approaches, and the biological rationale for GGT, bilirubin, and GBR."),
    ("Introduction", "3", "State specific objectives, including any prespecified hypotheses.", "Introduction", "The manuscript states the objective to examine the association between GBR and depressive symptoms, compare GBR with related indices, and evaluate heterogeneity by BMI, sex, and race/ethnicity."),
    ("Methods", "4", "Present key elements of study design early in the paper.", "Materials and methods: Study design and data source", "The study is described as a cross-sectional secondary analysis of NHANES 2007-2018."),
    ("Methods", "5", "Describe the setting, locations, and relevant dates, including periods of data collection.", "Materials and methods: Study design and data source", "The setting is the US NHANES survey, using 2007-2018 cycles; later cycles are excluded due to potential COVID-19-related operational changes."),
    ("Methods", "6", "Give the eligibility criteria, and the sources and methods of participant selection.", "Materials and methods: Study participants; S1 Fig", "Adults aged 18 years and older were selected; participants missing MEC weights were excluded; S1 Fig summarizes participant selection."),
    ("Methods", "7", "Clearly define outcomes, exposures, predictors, potential confounders, and effect modifiers.", "Materials and methods: Exposure, Outcome, Covariates, Statistical analysis", "The manuscript defines log(GBR), PHQ-9 total score and PHQ-9 >= 10, covariates, diabetes, CVD, medication use, subgroup variables, and interaction terms."),
    ("Methods", "8", "For each variable of interest, give sources of data and details of measurement methods.", "Materials and methods: Exposure, Outcome, Covariates", "Laboratory measures, PHQ-9 items, prescription medication files, questionnaire-based variables, and derived variables are described."),
    ("Methods", "9", "Describe any efforts to address potential sources of bias.", "Materials and methods: Statistical analysis; Sensitivity analyses; Limitations", "The analysis used MICE for missing data, survey weights, robust standard errors, extensive covariate adjustment, complete-case comparison, exclusions, liver-safe restriction, hsCRP sensitivity analysis, and subgroup analyses."),
    ("Methods", "10", "Explain how the study size was arrived at.", "Materials and methods: Study participants; Results: Sample characteristics", "The study size reflects all eligible NHANES adults from 2007-2018 with available MEC examination weights after applying the predefined inclusion/exclusion criteria; no a priori sample-size calculation was performed because this was a secondary analysis of public survey data."),
    ("Methods", "11", "Explain how quantitative variables were handled in the analyses, including any groupings.", "Materials and methods: Exposure, Outcome, Covariates, Statistical analysis", "The manuscript describes natural-log transformation, mean-centering, standardized comparisons, PHQ-9 as continuous and binary, BMI categories, cotinine strata, OBS tertiles, and RCS modeling."),
    ("Methods", "12a", "Describe all statistical methods, including those used to control for confounding.", "Materials and methods: Statistical analysis", "The manuscript describes WLS models, covariate adjustment, Rubin pooling, Barnard-Rubin correction, cluster-robust SEs, survey weight handling, VIF, and FDR correction."),
    ("Methods", "12b", "Describe methods used to examine subgroups and interactions.", "Materials and methods: Statistical analysis", "The manuscript describes BMI, sex, age, race/ethnicity, cotinine, and OBS subgroup analyses, plus interaction models for BMI, antidepressant use, smoking, OBS, and race/ethnicity."),
    ("Methods", "12c", "Explain how missing data were addressed.", "Materials and methods: Statistical analysis", "MICE with m = 20, IterativeImputer settings, inclusion of MEC weights as auxiliary variables, and CCA sensitivity analyses are described."),
    ("Methods", "12d", "Describe analytical methods taking account of the sampling strategy.", "Materials and methods: Statistical analysis; Limitations", "MEC weights divided by six cycles and PSU-based cluster-robust SEs were used; the limitations note incomplete accounting for the full NHANES stratification structure."),
    ("Methods", "12e", "Describe any sensitivity analyses.", "Materials and methods: Statistical analysis", "Sensitivity analyses include logistic PHQ-9 >= 10 models, complete-case analysis, exclusion of antidepressant users, CVD, diabetes, combined exclusions, ALT/AST-normal restriction, cotinine strata, and hsCRP-restricted models."),
    ("Results", "13a", "Report numbers of individuals at each stage of the study.", "Results: Sample characteristics; S1 Fig", "The final cohort size is reported as N = 36,580; S1 Fig provides the participant flow."),
    ("Results", "13b", "Give reasons for non-participation at each stage.", "Materials and methods: Study participants; S1 Fig", "The principal analytic exclusion was missing MEC examination weights; reasons for non-participation beyond NHANES public data availability were not directly available."),
    ("Results", "13c", "Consider use of a flow diagram.", "S1 Fig", "A participant selection flowchart is included as S1 Fig."),
    ("Results", "14a", "Give characteristics of study participants and information on exposures and potential confounders.", "Results: Sample characteristics; Table 1", "Table 1 reports participant characteristics by log(GBR) quartiles, including exposure components, demographic factors, lifestyle variables, comorbidities, and medication use."),
    ("Results", "14b", "Indicate the number of participants with missing data for each variable of interest.", "S7 Table; Methods; Results: Sensitivity analyses; S3 Table", "S7 Table reports variable-level availability and missingness in the analytic sample before MICE. Missing data handling is described through MICE and CCA comparison."),
    ("Results", "15", "Report numbers of outcome events or summary measures.", "Results: Primary analysis; Table 2", "PHQ-9 was analyzed as a continuous outcome, and PHQ-9 >= 10 was analyzed as a binary sensitivity outcome with ORs reported."),
    ("Results", "16a", "Give unadjusted and adjusted estimates and their precision; make clear which confounders were adjusted for.", "Results: Primary analysis; Table 2; Supplementary tables", "The fully adjusted beta estimate, SE, 95% CI, p-value, q-value, and logistic OR are reported; adjusted covariates are listed in Methods."),
    ("Results", "16b", "Report category boundaries when continuous variables are categorized.", "Methods and Results: Subgroup analyses", "BMI category boundaries, cotinine thresholds, PHQ-9 >= 10, and OBS tertile analyses are stated."),
    ("Results", "16c", "If relevant, consider translating relative risk estimates into absolute risk for a meaningful time period.", "Not applicable", "The primary outcome was continuous PHQ-9 in a cross-sectional analysis; temporal absolute risk translation was not applicable."),
    ("Results", "17", "Report other analyses, such as subgroups, interactions, and sensitivity analyses.", "Results: Sensitivity analyses, Subgroup analyses, Interaction analyses, Dose-response analyses, Comparison with other indices", "The manuscript reports sensitivity analyses, subgroup results, race/ethnicity interaction, antidepressant interaction, RCS analysis, specificity analyses, and OBS/CDAI comparisons."),
    ("Discussion", "18", "Summarize key results with reference to study objectives.", "Discussion: opening paragraph", "The discussion summarizes the independent association of GBR with PHQ-9 and highlights race/ethnicity-specific heterogeneity."),
    ("Discussion", "19", "Discuss limitations, considering potential sources of bias or imprecision.", "Discussion: Limitations", "The limitations section discusses cross-sectional design, missing data assumptions, unmeasured genetic/environmental confounding, single-time-point biomarkers, lack of bilirubin fractions, residual confounding, hsCRP cycle limitations, and incomplete survey-stratification handling."),
    ("Discussion", "20", "Give a cautious overall interpretation considering objectives, limitations, multiplicity, and evidence from similar studies.", "Discussion: Biological interpretation, Racial/ethnic specificity, Limitations, Conclusions", "The interpretation is cautious, frames findings as hypothesis-generating, notes modest effect size, and avoids clinical prediction claims."),
    ("Discussion", "21", "Discuss the generalisability of the study results.", "Discussion: Racial/ethnic specificity; Conclusions", "The manuscript explicitly states that GBR may not have population-invariant meaning and calls for external validation in Asian and Hispanic cohorts."),
    ("Other information", "22", "Give the source of funding and the role of funders.", "Funding", "The manuscript states that the author received no specific funding for this work."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")
        borders.append(element)
    tbl_pr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, *, size=9, bold=False, color="000000") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(cell, text: str, *, bold=False, size=8.5) -> None:
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    format_run(run, size=size, bold=bold)


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("STROBE Checklist for Cross-Sectional Studies")
    format_run(run, size=16, bold=True, color="1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(
        "Manuscript: Race- and ethnicity-specific association of the gamma-glutamyl transferase-to-bilirubin ratio with depressive symptoms in US adults: a NHANES 2007-2018 study"
    )
    format_run(run, size=9.5, bold=False, color="333333")

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run(
        "Checklist completed for a cross-sectional secondary analysis of publicly available NHANES data. Locations are given by manuscript section/table/figure because pagination may change after journal or preprint conversion."
    )
    format_run(run, size=9, color="333333")

    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    set_table_borders(table)
    headers = ["Section/topic", "Item", "STROBE recommendation", "Reported location", "Response / notes"]
    widths = [1550, 650, 3650, 2300, 4850]
    for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        set_cell_width(cell, width)
        set_cell_shading(cell, "E8EEF5")
        cell.text = ""
        add_paragraph(cell, header, bold=True, size=8.5)
    set_repeat_table_header(table.rows[0])

    for section_topic, item, recommendation, location, response in CHECKLIST_ROWS:
        cells = table.add_row().cells
        values = [section_topic, item, recommendation, location, response]
        for cell, value, width in zip(cells, values, widths):
            set_cell_width(cell, width)
            cell.text = ""
            add_paragraph(cell, value, bold=(item == "Item"), size=8.2)

    doc.add_paragraph()
    source = doc.add_paragraph()
    run = source.add_run(
        "Checklist source: STROBE Statement checklist for cross-sectional studies, STROBE Initiative (https://www.strobe-statement.org/checklists/)."
    )
    format_run(run, size=8.5, color="555555")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_markdown(path: Path) -> None:
    lines = [
        "# STROBE Checklist for Cross-Sectional Studies",
        "",
        "Manuscript: Race- and ethnicity-specific association of the gamma-glutamyl transferase-to-bilirubin ratio with depressive symptoms in US adults: a NHANES 2007-2018 study",
        "",
        "Locations are given by manuscript section/table/figure because pagination may change after journal or preprint conversion.",
        "",
        "| Section/topic | Item | STROBE recommendation | Reported location | Response / notes |",
        "|---|---:|---|---|---|",
    ]
    for row in CHECKLIST_ROWS:
        escaped = [value.replace("|", "\\|") for value in row]
        lines.append("| " + " | ".join(escaped) + " |")
    lines.extend(
        [
            "",
            "Checklist source: STROBE Statement checklist for cross-sectional studies, STROBE Initiative (https://www.strobe-statement.org/checklists/).",
            "",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    build_docx(OUT)
    build_docx(TP_OUT)
    build_markdown(MD_OUT)
    build_markdown(TP_MD_OUT)
    print(f"Saved {OUT}")
    print(f"Saved {TP_OUT}")
    print(f"Saved {MD_OUT}")
    print(f"Saved {TP_MD_OUT}")


if __name__ == "__main__":
    main()
