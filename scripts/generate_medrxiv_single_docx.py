import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from generate_tp_manuscript import (
    build_table_1,
    build_table_2,
    build_table_3,
    build_table_4,
    build_supp_table_1,
    build_supp_table_2,
    build_supp_table_3,
    build_supp_table_4,
    build_supp_table_5,
    build_supp_table_6,
    get_latest_results_file,
    apply_three_line_table_styling,
    apply_fonts_to_run,
)


ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "results"
OUT = ROOT / "submission_packages" / "medrxiv" / "medrxiv_single_manuscript.docx"
OUT_PRINT = ROOT / "submission_packages" / "medrxiv" / "medrxiv_single_manuscript_print_reading.docx"
GITHUB_URL = "https://github.com/mattakuya/NHANES_GBR"

MAJOR_HEADINGS = {
    "Abstract",
    "Introduction",
    "Materials and methods",
    "Results",
    "Discussion",
    "Conclusions",
    "Author contributions",
    "Funding",
    "Competing interests",
    "Use of generative AI and AI-assisted technologies",
    "Data availability",
    "References",
}

SUBHEADINGS = {
    "Study design and data source",
    "Ethical considerations",
    "Study participants",
    "Exposure",
    "Outcome",
    "Covariates",
    "Statistical analysis",
    "Sample characteristics",
    "Primary analysis",
    "Specificity analysis",
    "Sensitivity analyses",
    "Subgroup analyses",
    "Interaction analyses",
    "Dose-response analyses",
    "Comparison with other established oxidative stress indices",
    "Biological interpretation of GBR",
    "GBR relative to GGT alone",
    "Independence and complementarity relative to existing oxidative balance indices",
    "Heterogeneity of effects by BMI",
    "Racial/ethnic specificity",
    "Interaction between GBR and antidepressant use",
    "Effect size and epidemiological relevance",
    "Limitations",
}


def clean_inline(text: str) -> str:
    text = re.sub(r"_([^_\s][^_]*[^_\s])_", r"\1", text)
    return text.replace("**", "").replace("`", "").replace("&ge;", "≥").replace("&le;", "≤")


def set_margins(section, top=0.75, bottom=0.75, left=0.75, right=0.75):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    set_margins(section, top=0.7, bottom=0.7, left=0.7, right=0.7)


def add_page_number_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_paragraph(doc, text, size=11, bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    run = p.add_run(clean_inline(text))
    apply_fonts_to_run(run, bold=bold, size_pt=size)
    run.italic = italic
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = add_paragraph(doc, text, size=13, bold=True, before=14, after=6)
        p.paragraph_format.keep_with_next = True
        return p
    if level == 2:
        p = add_paragraph(doc, text, size=11.5, bold=True, before=10, after=4)
        p.paragraph_format.keep_with_next = True
        return p
    p = add_paragraph(doc, text, size=10.5, bold=True, before=7, after=3)
    p.paragraph_format.keep_with_next = True
    return p


def markdown_table_rows(md_table: str):
    rows = []
    for line in md_table.strip().splitlines():
        if not line.strip().startswith("|"):
            continue
        if re.match(r"^\s*\|\s*[:\-|\s]+$", line):
            continue
        rows.append([clean_inline(c.strip()) for c in line.strip().split("|")[1:-1]])
    return rows


def add_markdown_table(doc, title, md_table, note=None, landscape=False):
    doc.add_page_break()
    if landscape:
        set_landscape(doc.add_section())
    else:
        section = doc.add_section()
        set_margins(section)
    add_heading(doc, title, level=2)
    if note:
        add_paragraph(doc, note, size=9.5, italic=True, after=5)
    rows = markdown_table_rows(md_table)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = None
    apply_three_line_table_styling(table, rows)


def add_inline_markdown_table(doc, md_table, note=None):
    if note:
        add_paragraph(doc, note, size=9.5, italic=True, after=5)
    rows = markdown_table_rows(md_table)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = None
    apply_three_line_table_styling(table, rows)
    add_paragraph(doc, "", after=2)


def build_supp_table_7():
    path = ROOT / "submission_packages" / "medrxiv" / "supp_table7_missingness.md"
    if not path.exists():
        raise FileNotFoundError(
            "Missingness table not found. Run generate_missingness_table.py first."
        )
    return path.read_text(encoding="utf-8")


def add_inline_figure(doc, image_path, caption=None, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    if caption:
        add_paragraph(doc, caption, size=9.5, italic=True, before=3, after=8)


def add_markdown_table_page(doc, title, md_table, note=None, landscape=False):
    doc.add_page_break()
    if landscape:
        set_landscape(doc.add_section())
    else:
        section = doc.add_section()
        set_margins(section)
    add_heading(doc, title, level=2)
    if note:
        add_paragraph(doc, note, size=9.5, italic=True, after=5)
    rows = markdown_table_rows(md_table)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = None
    apply_three_line_table_styling(table, rows)


def add_figure(doc, title, image_path, caption, width=6.4):
    doc.add_page_break()
    section = doc.add_section()
    set_margins(section)
    add_heading(doc, title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_paragraph(doc, caption, size=9.5, italic=True, before=4, after=8)


def manuscript_lines_for_single_doc():
    text = (ROOT / "submission_packages" / "translational_psychiatry" / "manuscript_tp.md").read_text(encoding="utf-8")
    text = text.replace("[repository URL]", GITHUB_URL)
    text = text.replace("\n[Insert Table 1 here]\n", "\n")
    text = text.replace("\n[Insert Table 2 here]\n", "\n")
    text = text.replace("\n[Insert Table 3 here]\n", "\n")
    text = text.replace("\n[Insert Table 4 here]\n", "\n")
    # Keep the supporting-information index out of the main text; the full supplement is appended.
    text = text.replace("\nSupporting information\n", "\n")
    text = re.sub(r"\n\*\*S1 Checklist\.[\s\S]*?(?=\nReferences\n)", "\n", text)
    return text.splitlines()


def add_markdown_body(doc, inline_display_items=False, results_file=None):
    inserted = set()
    title_done = False
    for raw in manuscript_lines_for_single_doc():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#### Table 3."):
            if inline_display_items:
                add_paragraph(doc, line[5:], size=10.5, bold=True, before=5, after=4)
                add_inline_markdown_table(doc, build_table_3(results_file))
                inserted.add("table3")
            continue
        if line.startswith("#### Table 4."):
            if inline_display_items:
                add_paragraph(doc, line[5:], size=10.5, bold=True, before=5, after=4)
                add_inline_markdown_table(doc, build_table_4(results_file))
                inserted.add("table4")
            continue
        if line in ["[Insert Table 3 here]", "[Insert Table 4 here]"]:
            continue
        if not title_done:
            p = add_paragraph(
                doc,
                line,
                size=16,
                bold=True,
                before=0,
                after=8,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            p.paragraph_format.keep_with_next = True
            title_done = True
            continue
        if line in MAJOR_HEADINGS:
            add_heading(doc, line, level=1)
            continue
        if line in SUBHEADINGS:
            add_heading(doc, line, level=2)
            continue
        if inline_display_items:
            if line.startswith("**Table 1."):
                add_paragraph(doc, line, size=10.5, bold=True, before=5, after=4)
                add_inline_markdown_table(
                    doc,
                    build_table_1(results_file),
                    note="Values are unweighted counts, means (SD), or weighted percentages, as indicated. GBR, gamma-glutamyl transferase-to-total bilirubin ratio.",
                )
                inserted.add("table1")
                continue
            if line.startswith("**Table 2."):
                add_paragraph(doc, line, size=10.5, bold=True, before=5, after=4)
                add_inline_markdown_table(
                    doc,
                    build_table_2(results_file),
                    note="Estimates are from fully adjusted weighted models unless otherwise indicated. The logistic model reports odds ratios for PHQ-9 ≥ 10.",
                )
                inserted.add("table2")
                continue
            if line.startswith("**Fig. 1."):
                add_paragraph(doc, line, size=10.5, bold=True, before=5, after=4)
                add_inline_figure(
                    doc,
                    RESULTS / "tp_figures" / "figure1_race_ethnicity_interaction.png",
                    width=6.4,
                )
                inserted.add("fig1")
                continue
            if line.startswith("**Fig. 2."):
                add_paragraph(doc, line, size=10.5, bold=True, before=5, after=4)
                add_inline_figure(
                    doc,
                    RESULTS / "tp_figures" / "figure2_specificity.png",
                    width=6.2,
                )
                inserted.add("fig2")
                continue
            if line.startswith("**Fig. 3."):
                add_paragraph(doc, line, size=10.5, bold=True, before=5, after=4)
                add_inline_figure(
                    doc,
                    RESULTS / "tp_figures" / "figure3_rcs.png",
                    width=6.2,
                )
                inserted.add("fig3")
                continue
        if line.startswith("# "):
            add_paragraph(doc, line[2:], size=16, bold=True, before=4, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif line.startswith("## "):
            add_heading(doc, line[3:], level=1)
        elif line.startswith("### "):
            add_heading(doc, line[4:], level=2)
        elif line.startswith("#### "):
            add_heading(doc, line[5:], level=3)
        elif line.startswith("**") and line.endswith("**"):
            add_paragraph(doc, line, size=10.5, bold=True, before=5, after=4)
        else:
            add_paragraph(doc, line)
    return inserted


def main():
    results_file = get_latest_results_file()
    race_results_file = results_file

    doc = Document()
    set_margins(doc.sections[0])
    add_page_number_footer(doc)

    add_markdown_body(doc)

    add_markdown_table(
        doc,
        "Table 1. Participant characteristics by quartiles of log(GBR).",
        build_table_1(results_file),
        note="Values are unweighted counts, means (SD), or weighted percentages, as indicated. GBR, gamma-glutamyl transferase-to-total bilirubin ratio.",
    )
    add_markdown_table(
        doc,
        "Table 2. Primary and sensitivity analyses of the association between log(GBR) and depressive symptoms.",
        build_table_2(results_file),
        note="Estimates are from fully adjusted weighted models unless otherwise indicated. The logistic model reports odds ratios for PHQ-9 ≥ 10.",
    )
    add_markdown_table(
        doc,
        "Table 3. Subgroup-stratified analyses of the association between log(GBR) and depressive symptoms (WLS models).",
        build_table_3(results_file),
        landscape=False,
    )
    add_markdown_table(
        doc,
        "Table 4. Independent and additive associations of GBR and OBS (standardized per 1 SD).",
        build_table_4(results_file),
        landscape=False,
    )

    add_figure(
        doc,
        "Figure 1. Race/ethnicity-specific association of log(GBR) with depressive symptoms.",
        RESULTS / "tp_figures" / "figure1_race_ethnicity_interaction.png",
        "a, Race/ethnicity-stratified adjusted beta coefficients for PHQ-9 total score. b, Interaction contrasts relative to Non-Hispanic White participants. Error bars denote 95% confidence intervals.",
        width=6.7,
    )
    add_figure(
        doc,
        "Figure 2. Specificity of the GBR association relative to liver enzymes and bilirubin.",
        RESULTS / "tp_figures" / "figure2_specificity.png",
        "Standardized beta coefficients and 95% confidence intervals are shown for log(GBR), log(GGT), log(total bilirubin), ALT, and AST in relation to PHQ-9 total score.",
        width=6.4,
    )
    add_figure(
        doc,
        "Figure 3. Dose-response association between log(GBR) and depressive symptoms.",
        RESULTS / "tp_figures" / "figure3_rcs.png",
        "Restricted cubic spline model with four knots showing the adjusted association between log(GBR) and PHQ-9 total score. Shading denotes the 95% confidence interval; vertical dashed lines denote knot positions.",
        width=6.4,
    )

    doc.add_page_break()
    section = doc.add_section()
    set_margins(section)
    add_heading(doc, "Supplementary material", level=1)
    add_paragraph(doc, "The following supplementary tables and figures are included for completeness in this single-file preprint package.", size=10.5)

    supp_tables = [
        ("Supplementary Table S1. Multicollinearity diagnostics for main model covariates.", build_supp_table_1(results_file)),
        ("Supplementary Table S2. Model fit comparison of the primary exposure variables.", build_supp_table_2(results_file)),
        ("Supplementary Table S3. Complete case analysis versus multiple imputation by chained equations (MICE).", build_supp_table_3(results_file)),
        ("Supplementary Table S4. Associations with established oxidative stress indicators and joint models.", build_supp_table_4(results_file)),
        ("Supplementary Table S5. Antidepressant interaction analyses.", build_supp_table_5(results_file)),
        ("Supplementary Table S6. Race/ethnicity interaction analysis.", build_supp_table_6(str(race_results_file))),
        ("Supplementary Table S7. Missingness of primary analysis variables before multiple imputation.", build_supp_table_7()),
    ]
    for title, table_md in supp_tables:
        add_markdown_table(doc, title, table_md, landscape=False)

    supp_figs = [
        ("Supplementary Fig. S1. Participant selection flowchart.", "supp_figure1_flowchart.png", "Selection of eligible adults from NHANES 2007-2018.", 6.0),
        ("Supplementary Fig. S2. Demographic subgroup analyses.", "supp_figure2_demographic_subgroups.png", "Adjusted beta coefficients and 95% confidence intervals across BMI, sex, and age strata.", 6.4),
        ("Supplementary Fig. S3. Lifestyle and clinical subgroup analyses.", "supp_figure3_lifestyle_clinical_subgroups.png", "Adjusted beta coefficients and 95% confidence intervals across alcohol, serum cotinine, and hsCRP-related analyses.", 6.4),
        ("Supplementary Fig. S4. Antidepressant use interaction.", "supp_figure4_ad_interaction.png", "Predicted PHQ-9 total score by log(GBR), stratified by antidepressant use.", 6.2),
    ]
    for title, filename, caption, width in supp_figs:
        add_figure(doc, title, RESULTS / "tp_figures" / filename, caption, width=width)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved {OUT}")


def main_print_reading():
    results_file = get_latest_results_file()
    race_results_file = results_file

    doc = Document()
    set_margins(doc.sections[0])
    add_page_number_footer(doc)

    inserted = add_markdown_body(doc, inline_display_items=True, results_file=results_file)

    doc.add_page_break()
    section = doc.add_section()
    set_margins(section)
    add_heading(doc, "Supplementary material", level=1)
    add_paragraph(doc, "Supplementary tables and figures are appended after the main text for print-reading convenience.", size=10.5)

    supp_tables = [
        ("Supplementary Table S1. Multicollinearity diagnostics for main model covariates.", build_supp_table_1(results_file)),
        ("Supplementary Table S2. Model fit comparison of the primary exposure variables.", build_supp_table_2(results_file)),
        ("Supplementary Table S3. Complete case analysis versus multiple imputation by chained equations (MICE).", build_supp_table_3(results_file)),
        ("Supplementary Table S4. Associations with established oxidative stress indicators and joint models.", build_supp_table_4(results_file)),
        ("Supplementary Table S5. Antidepressant interaction analyses.", build_supp_table_5(results_file)),
        ("Supplementary Table S6. Race/ethnicity interaction analysis.", build_supp_table_6(str(race_results_file))),
        ("Supplementary Table S7. Missingness of primary analysis variables before multiple imputation.", build_supp_table_7()),
    ]
    for title, table_md in supp_tables:
        add_markdown_table_page(doc, title, table_md, landscape=False)

    supp_figs = [
        ("Supplementary Fig. S1. Participant selection flowchart.", "supp_figure1_flowchart.png", "Selection of eligible adults from NHANES 2007-2018.", 6.0),
        ("Supplementary Fig. S2. Demographic subgroup analyses.", "supp_figure2_demographic_subgroups.png", "Adjusted beta coefficients and 95% confidence intervals across BMI, sex, and age strata.", 6.4),
        ("Supplementary Fig. S3. Lifestyle and clinical subgroup analyses.", "supp_figure3_lifestyle_clinical_subgroups.png", "Adjusted beta coefficients and 95% confidence intervals across alcohol, serum cotinine, and hsCRP-related analyses.", 6.4),
        ("Supplementary Fig. S4. Antidepressant use interaction.", "supp_figure4_ad_interaction.png", "Predicted PHQ-9 total score by log(GBR), stratified by antidepressant use.", 6.2),
    ]
    for title, filename, caption, width in supp_figs:
        add_figure(doc, title, RESULTS / "tp_figures" / filename, caption, width=width)

    OUT_PRINT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PRINT)
    print(f"Saved {OUT_PRINT}")
    print(f"Inline display items inserted: {sorted(inserted)}")


if __name__ == "__main__":
    main()
    main_print_reading()
