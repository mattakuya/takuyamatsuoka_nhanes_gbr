import os
import glob
import re
import math
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def get_latest_results_file(results_dir="results"):
    pattern = os.path.join(results_dir, "nhanes_results_*.md")
    files = glob.glob(pattern)
    # Exclude QUICK files
    files = [f for f in files if "QUICK" not in f]
    if not files:
        # Check parent directory just in case
        pattern_parent = os.path.join("..", results_dir, "nhanes_results_*.md")
        files = glob.glob(pattern_parent)
        files = [f for f in files if "QUICK" not in f]
        if not files:
            raise FileNotFoundError("Results file matching 'nhanes_results_*.md' not found in results directory.")
    return max(files, key=os.path.getmtime)

def parse_markdown_table(file_path, header_title):
    """
    Parses a markdown table from the results file following a specific heading title.
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Locate the heading
    pattern = rf"##\s+{re.escape(header_title)}\s*\n\n_FDR family:.*_\n\n(\|[\s\S]*?)(?=\n\n##|\Z)"
    match = re.search(pattern, content)
    if not match:
        # Try without the FDR family note
        pattern_simple = rf"##\s+{re.escape(header_title)}\s*\n\n(\|[\s\S]*?)(?=\n\n##|\Z)"
        match = re.search(pattern_simple, content)
    
    if match:
        return match.group(1).strip()
    else:
        raise ValueError(f"Could not find table under heading: {header_title}")

def build_table_1(results_file):
    return parse_markdown_table(results_file, "DESCRIPTIVE TABLE 1 BY LOG_RATIO QUARTILE")

def build_table_2(results_file):
    """
    Constructs Table 2 (Primary and Sensitivity Analyses in the Full Sample)
    """
    main_table = parse_markdown_table(results_file, "MAIN ANALYSIS: PHQ-9 CONTINUOUS")
    logistic_table = parse_markdown_table(results_file, "SENSITIVITY: PHQ-9 ≥ 10 (Logistic, β on log-odds)")
    sensitivity_table = parse_markdown_table(results_file, "SENSITIVITY ANALYSIS")
    liver_safe_table = parse_markdown_table(results_file, "LIVER-SAFE (NO DILI) ANALYSIS")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Analysis Model / Target | N | Beta / OR [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    
    # 1. Primary WLS
    rows.append("| **Primary Analysis (WLS)** | | | | | |")
    main_row = extract_rows(main_table)[0].split('|')[1:-1]
    subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = [x.strip() for x in main_row]
    rows.append(f"| Full Sample (OLS PHQ-9) | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
    
    # 2. Logistic PHQ-9 >= 10
    rows.append("| **Sensitivity Analysis (Logistic)** | | | | | |")
    log_row = extract_rows(logistic_table)[0].split('|')[1:-1]
    subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = [x.strip() for x in log_row]
    or_val = math.exp(float(beta))
    or_lo = math.exp(float(ci_lo))
    or_hi = math.exp(float(ci_hi))
    rows.append(f"| Full Sample (Logistic PHQ-9 &ge; 10) | {n} | OR = {or_val:.4f} [{or_lo:.4f}, {or_hi:.4f}] | — | {p} | {q} |")
    
    # 3. Exclusion sensitivity
    rows.append("| **Sensitivity Analyses (Exclusion Models)** | | | | | |")
    for row_str in extract_rows(sensitivity_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    # 4. Liver-safe
    rows.append("| **Liver-Safe Analysis (No DILI)** | | | | | |")
    for row_str in extract_rows(liver_safe_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def build_table_3(results_file):
    """
    Constructs Table 3 (Comprehensive Subgroup Analyses)
    """
    bmi_table = parse_markdown_table(results_file, "BMI STRATIFIED ANALYSIS")
    sex_table = parse_markdown_table(results_file, "SEX-STRATIFIED ANALYSIS")
    age_table = parse_markdown_table(results_file, "AGE-STRATIFIED ANALYSIS")
    race_table = parse_markdown_table(results_file, "RACE-STRATIFIED ANALYSIS")
    alcohol_table = parse_markdown_table(results_file, "ALCOHOL-STRATIFIED ANALYSIS (NIAAA)")
    cotinine_table = parse_markdown_table(results_file, "COTININE-STRATIFIED ANALYSIS (Pirkle)")
    crp_table = parse_markdown_table(results_file, "CRP-STRATIFIED ANALYSIS (AHA, I/J only)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Subgroup / Stratification | N | Beta [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    
    def append_group(title, table_str, skip_full=True):
        rows.append(f"| **{title}** | | | | | |")
        for row_str in extract_rows(table_str):
            parts = [x.strip() for x in row_str.split('|')[1:-1]]
            subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
            if skip_full and "Full Sample" in subgroup:
                continue
            rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
            
    append_group("BMI Categories (kg/m²)", bmi_table)
    append_group("Sex", sex_table)
    append_group("Age Groups", age_table)
    append_group("Race / Ethnicity", race_table)
    append_group("Alcohol Consumption (NIAAA)", alcohol_table)
    append_group("Cotinine Level (Smoking Exposure)", cotinine_table)
    append_group("Systemic Inflammation (hs-CRP, I/J only)", crp_table)
    
    return '\n'.join([header, separator] + rows)

def build_table_4(results_file):
    """
    Constructs Table 4 (Specificity Analyses standardized per 1 SD)
    """
    spec_table = parse_markdown_table(results_file, "SPECIFICITY ANALYSIS (SD-STANDARDIZED)")
    joint_table = parse_markdown_table(results_file, "BILIRUBIN BEYOND GGT (JOINT MODEL, per 1 SD)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Model / Term | N | Beta [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    
    # Univariate
    rows.append("| **Univariate Standardized Models** | | | | | |")
    for row_str in extract_rows(spec_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    # Joint
    rows.append("| **Joint Model of GGT + Bilirubin** | | | | | |")
    for row_str in extract_rows(joint_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        term, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def build_table_5(results_file):
    """
    Constructs Table 5 (Independent and Additive Association of GBR and OBS per 1 SD)
    """
    obs_table = parse_markdown_table(results_file, "JOINT MODEL: log_ratio + OBS (per 1 SD each, SD-standardized)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Subgroup / Model | Term | N | Beta [95% CI] | SE | P-value | Attenuation (%) |"
    separator = "| :--- | :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    for row_str in extract_rows(obs_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        if len(parts) == 13:
            subgroup, model, term_part1, term_part2, n, beta, se, ci_lo, ci_hi, p, atten, q, q_fam = parts
            term = f"{term_part1} + {term_part2}"
        else:
            subgroup, model, term, n, beta, se, ci_lo, ci_hi, p, atten, q, q_fam = parts
        rows.append(f"| {subgroup} ({model}) | {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {atten} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_1(results_file):
    return parse_markdown_table(results_file, "VIF (Main Model Covariates)")

def build_supp_table_2(results_file):
    model_fit_table = parse_markdown_table(results_file, "MODEL FIT COMPARISON: RATIO vs GGT vs GGT+BIL")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Model | N | Parameters | Pseudo-AIC | Pseudo-BIC | Adjusted R² | Weighted RMSE |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    for row_str in extract_rows(model_fit_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        model, n, m, aic_mean, aic_sd, bic_mean, bic_sd, r2_mean, r2_sd, rmse_mean, rmse_sd = parts
        rows.append(f"| {model} | {n} | {m} | {float(aic_mean):.2f} | {float(bic_mean):.2f} | {float(r2_mean):.4f} | {float(rmse_mean):.4f} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_3(results_file):
    return parse_markdown_table(results_file, "CCA vs MICE: GGT BIAS ASSESSMENT (SD-standardized)")

def build_supp_table_4(results_file):
    spec_table = parse_markdown_table(results_file, "OXIDATIVE COMPARISON: GGT/Bil vs CDAI vs OBS (SD-STANDARDIZED)")
    excl_ad_table = parse_markdown_table(results_file, "OXIDATIVE COMPARISON EXCL. AD USERS (SD-STANDARDIZED)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Analysis Cohort / Indicator | N | Beta [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    rows.append("| **Full Cohort (OBS-Restricted Sample)** | | | | | |")
    for row_str in extract_rows(spec_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    rows.append("| **Antidepressant Non-Users Cohort** | | | | | |")
    for row_str in extract_rows(excl_ad_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        subgroup, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {subgroup} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_5(results_file):
    ad_table = parse_markdown_table(results_file, "AD-INTERACTION ANALYSIS [log_ratio]")
    ad_normal_table = parse_markdown_table(results_file, "AD-INTERACTION ANALYSIS [log_ratio] (Normal BMI)")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Sample / Term | N | Beta [95% CI] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    rows.append("| **Full Cohort Interaction Model** | | | | | |")
    for row_str in extract_rows(ad_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        term, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    rows.append("| **Normal BMI (18.5–24.9 kg/m²) Cohort** | | | | | |")
    for row_str in extract_rows(ad_normal_table):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        term, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def build_supp_table_6(race_results_file):
    race_interaction = parse_markdown_table(race_results_file, "RACE/ETHNICITY INTERACTION ANALYSIS [log_ratio]")
    global_test = parse_markdown_table(race_results_file, "GLOBAL RACE/ETHNICITY INTERACTION TEST [log_ratio]")
    
    def extract_rows(table_str):
        return table_str.strip().split('\n')[2:]
        
    header = "| Term / Test | N | Beta / Chi-square [95% CI / df] | SE | P-value | Q-value |"
    separator = "| :--- | :---: | :---: | :---: | :---: | :---: |"
    
    rows = []
    rows.append("| **Interaction Model (NHWhite as Reference)** | | | | | |")
    for row_str in extract_rows(race_interaction):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        term, n, beta, se, ci_lo, ci_hi, p, q, q_fam = parts
        rows.append(f"| {term} | {n} | {beta} [{ci_lo}, {ci_hi}] | {se} | {p} | {q} |")
        
    rows.append("| **Global Interaction Wald Test** | | | | | |")
    for row_str in extract_rows(global_test):
        parts = [x.strip() for x in row_str.split('|')[1:-1]]
        test, n, chi2, df, p, terms, q, q_fam = parts
        rows.append(f"| {test} | {n} | Chi² = {float(chi2):.4f} [df = {df}] | — | {p} | {q} |")
        
    return '\n'.join([header, separator] + rows)

def parse_estimate(val_str):
    # Parses beta, ci_lo, ci_hi from e.g. "0.3013 [0.1541, 0.4484]"
    match = re.search(r"(-?\d+\.\d+)\s*\[\s*(-?\d+\.\d+)\s*,\s*(-?\d+\.\d+)\s*\]", val_str)
    if match:
        return float(match.group(1)), float(match.group(2)), float(match.group(3))
    return None

def format_p_value(p_str):
    try:
        p_val = float(p_str)
        if p_val < 0.001:
            return "< 0.001"
        else:
            return f"{p_val:.3f}"
    except:
        return p_str

def markdown_to_dataframe(md_table_str):
    lines = md_table_str.strip().split('\n')
    # Filter out separator lines like | :--- | :---: |
    lines = [l for l in lines if not re.match(r'^\s*\|?\s*:?-+:?\s*\|', l)]
    
    if not lines:
        return pd.DataFrame()
        
    header_line = lines[0]
    headers = [x.strip() for x in header_line.split('|')[1:-1]]
    
    rows = []
    for line in lines[1:]:
        if not line.strip():
            continue
        row = [x.strip() for x in line.split('|')[1:-1]]
        if len(row) < len(headers):
            row += [''] * (len(headers) - len(row))
        elif len(row) > len(headers):
            row = row[:len(headers)]
        rows.append(row)
        
    return pd.DataFrame(rows, columns=headers)

def apply_fonts_to_run(run, bold=False, size_pt=10):
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Helvetica')
    rFonts.set(qn('w:hAnsi'), 'Helvetica')
    rPr.append(rFonts)
    
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True

def apply_three_line_table_styling(table, rows_data):
    """
    Applies professional booktabs (three-line table) styling to a python-docx table.
    """
    tblPr = table._tbl.tblPr
    
    # 1. Clear standard borders by adding custom w:tblBorders
    tblBorders = OxmlElement('w:tblBorders')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '111111')
    tblBorders.append(top)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'none')
    tblBorders.append(bottom)
    
    for border_name in ('left', 'right', 'insideV', 'insideH'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
        
    tblPr.append(tblBorders)
    
    # 2. Set table-level cell padding (dxa unit)
    tblCellMar = OxmlElement('w:tblCellMar')
    for margin_name, val in [('top', 100), ('bottom', 100), ('left', 150), ('right', 150)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)
    tblPr.append(tblCellMar)
    
    # 3. Render and style cells
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        
        # Check if this is a group header row (all bold or empty other cells)
        is_group_header = False
        first_val = row_data[0].strip()
        if r_idx > 0 and (first_val.startswith('**') and first_val.endswith('**')):
            is_group_header = True
        elif r_idx > 0 and all(cell.strip() == '' for cell in row_data[1:]):
            is_group_header = True
            
        if is_group_header:
            # Merge all cells in this row
            merged_cell = row.cells[0]
            for c_idx in range(1, len(row.cells)):
                merged_cell = merged_cell.merge(row.cells[c_idx])
            
            # Style the merged group header cell
            merged_cell.text = ""
            p = merged_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(first_val.replace('**', ''))
            apply_fonts_to_run(run, bold=True, size_pt=10)
            
            # Subtle background color (#F2F2F2)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'F2F2F2')
            merged_cell._tc.get_or_add_tcPr().append(shading)
            continue
            
        for c_idx, cell_value in enumerate(row_data):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                # Alignments
                if c_idx == 0:
                    p.alignment = 0  # Left
                else:
                    p.alignment = 1  # Center
                
                is_header = (r_idx == 0)
                clean_val = cell_value.replace("**", "").replace("&ge;", "≥").replace("&le;", "≤")
                
                run = p.add_run(clean_val)
                apply_fonts_to_run(run, bold=is_header, size_pt=10)
                
                # Add thin bottom border to header row (row 0)
                if r_idx == 0:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')  # 0.75 pt
                    bottom.set(qn('w:space'), '0')
                    bottom.set(qn('w:color'), '111111')
                    tcBorders.append(bottom)
                    tcPr.append(tcBorders)

                # Add thick bottom border to last row (bottom rule)
                if r_idx == len(rows_data) - 1:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '12')  # 1.5 pt
                    bottom.set(qn('w:space'), '0')
                    bottom.set(qn('w:color'), '111111')
                    tcBorders.append(bottom)
                    tcPr.append(tcBorders)

    # 4. Prevent rows from splitting across pages
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    doc = Document()
    
    # Margin settings (1 inch = Inches(1))
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
        
    in_table = False
    table_lines = []
    
    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        
        rows_data = []
        for line in table_lines:
            if not line.strip() or re.match(r'^\s*\|\s*[:\-|\s]+$', line):
                continue
            parts = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            rows_data.append(parts)
            
        if not rows_data:
            table_lines = []
            return
            
        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = None
        apply_three_line_table_styling(table, rows_data)
        
        doc.add_paragraph()
        table_lines = []
        
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.strip().startswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False
            
        if not line.strip():
            i += 1
            continue
            
        # Heading 1
        h1_match = re.match(r'^#\s+(.*)', line)
        if h1_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(h1_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=18)
            i += 1
            continue
            
        # Heading 2
        h2_match = re.match(r'^##\s+(.*)', line)
        if h2_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h2_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=14)
            i += 1
            continue
            
        # Heading 3
        h3_match = re.match(r'^###\s+(.*)', line)
        if h3_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h3_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=12)
            i += 1
            continue
            
        # Heading 4
        h4_match = re.match(r'^####\s+(.*)', line)
        if h4_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h4_match.group(1))
            apply_fonts_to_run(run, bold=True, size_pt=11)
            i += 1
            continue
            
        # Normal Paragraph with bold parsing
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                bold_text = part[2:-2]
                run = p.add_run(bold_text)
                apply_fonts_to_run(run, bold=True, size_pt=11)
            else:
                run = p.add_run(part)
                apply_fonts_to_run(run, bold=False, size_pt=11)
                
        i += 1
        
    if in_table:
        flush_table()
        
    doc.save(docx_path)
    print(f"Successfully saved Word file to: {docx_path}")
